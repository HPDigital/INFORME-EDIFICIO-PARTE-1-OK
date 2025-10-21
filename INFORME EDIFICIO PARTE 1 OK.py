"""
INFORME EDIFICIO PARTE 1 OK
"""

#!/usr/bin/env python
# coding: utf-8

# In[2]:


import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os

# Lista de departamentos en el orden específico
orden_departamentos = [
    '1A', '1B', '1C', '1D', '1E',
    '2A', '2B', '2C', '2D', '2E',
    '3A', '3B', '3C', '3D', '3E',
    '4A', '4B', '4C', '4D', '4E',
    '5A', '5B', '5C', '5D', '5E',
    '6A', '6B', '6C', '6D', '6E',
    '7A', '7B', '7C', '7D', '7E',
    '8A', '8B', '8C', '8D', '8E',
    '9A', '9B', '9C', '9D', '9E',
    'T1', 'T2', 'T3', 'T4', 'T5']

def cargar_datos_excel(ruta_archivo, hoja):
    """Función para cargar datos desde un archivo Excel."""
    try:
        with pd.ExcelFile(ruta_archivo) as xls:
            return pd.read_excel(xls, sheet_name=hoja)
    except FileNotFoundError:
        print(f"El archivo no se encontró en la ruta: {ruta_archivo}")
        return None

def filtrar_departamentos_en_mora(df):
    """Función para filtrar departamentos que están en mora."""
    return df[df['MORA'] == 'ESTA EN MORA']

def generar_grafico_barras_con_saldo(df, columna_departamento, columna_saldo, nombre_imagen, orden_departamentos):
    """Función para generar un gráfico de barras horizontales con colores según el saldo."""

    # Asegurarse de que los departamentos están en el orden correcto
    df[columna_departamento] = pd.Categorical(df[columna_departamento], categories=orden_departamentos, ordered=True)
    df = df.sort_values(by=columna_departamento)

    plt.figure(figsize=(10, 13))

    # Definir colores: rojo para negativo y azul para positivo
    colores = df[columna_saldo].apply(lambda x: 'red' if x < 0 else 'blue')

    # Graficar los saldos con colores basados en su valor
    plt.barh(df[columna_departamento], df[columna_saldo], color=colores)

    plt.xlabel('Saldo')
    plt.ylabel(columna_departamento)
    plt.title('Conciliación Bancaria - Saldos por Departamento')

    # Agregar etiquetas de monto en las barras
    for i, v in enumerate(df[columna_saldo]):
        plt.text(v - 20 if v < 0 else v + 20, 
                 i, 
                 f'{v:,.2f}', 
                 va='center', 
                 ha='right' if v < 0 else 'left')

    plt.tight_layout()

    # Guardar el gráfico como imagen
    plt.savefig(nombre_imagen)
    plt.close()

    return nombre_imagen

def crear_informe_word(ruta_imagen_mora, ruta_imagen_conciliacion, ruta_salida):
    """Función para crear un documento Word e insertar los gráficos."""
    doc = Document()
    doc.add_heading('DEPARTAMENTOS/TIENDAS EN MORA', 0)

    # Insertar el gráfico de mora en el documento Word
    doc.add_picture(ruta_imagen_mora, width=Inches(6))  # Ajustar el tamaño del gráfico

    # Insertar el gráfico de conciliación en el documento Word
    doc.add_heading('Datos de Conciliación', level=1)
    doc.add_picture(ruta_imagen_conciliacion, width=Inches(6))  # Ajustar el tamaño del gráfico

    # Guardar el documento Word
    doc.save(ruta_salida)
    print(f"Informe generado y guardado en: {ruta_salida}")

# Ruta del archivo y hoja específica
file_path = r"C:\Users\HP\Desktop\EDIFICIO JUAN BOSCO\EXPENSAS\EXPENSAS CON CONCILIACION BANCARIA\INFOMRE DE ADEUDOS PYTHON.xlsx"
sheet_name = "EDIFICIO"

# Obtener la carpeta donde está el archivo Excel
carpeta_archivo = os.path.dirname(file_path)
ruta_salida_word = os.path.join(carpeta_archivo, "informe_departamentos_mora_conciliacion.docx")

# Ejecutar el flujo
df = cargar_datos_excel(file_path, sheet_name)

if df is not None:
    df_mora = filtrar_departamentos_en_mora(df)
    if not df_mora.empty:
        # Generar gráfico de departamentos en mora
        ruta_imagen_mora = generar_grafico_barras_con_saldo(
            df_mora, 'DEPARTAMENTO', 'SALDO',
            'grafico_mora.png', orden_departamentos
        )

        # Asegurarse de que la columna 'SALDO' esté disponible
        if 'SALDO' in df.columns:
            # Generar gráfico con barras horizontales según el saldo y en el orden especificado
            ruta_imagen_conciliacion = generar_grafico_barras_con_saldo(
                df, 'DEPARTAMENTO', 'SALDO', 'grafico_conciliacion_saldo.png', orden_departamentos
            )

            # Crear el informe de Word con ambos gráficos
            crear_informe_word(ruta_imagen_mora, ruta_imagen_conciliacion, ruta_salida_word)

            # Eliminar las imágenes temporales después de generar el informe
            if os.path.exists(ruta_imagen_mora):
                os.remove(ruta_imagen_mora)
            if os.path.exists(ruta_imagen_conciliacion):
                os.remove(ruta_imagen_conciliacion)
        else:
            print("La columna 'SALDO' no existe en el archivo Excel.")
    else:
        print("No se encontraron departamentos en mora.")


# In[ ]:






if __name__ == "__main__":
    pass
